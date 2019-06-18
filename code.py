from sys import argv
from pathlib import Path
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt, RGBColor, Mm
import io

if __name__ == '__main__':

    def styles(document):
        styles = document.styles
        new_heading_style = styles.add_style('filename', WD_STYLE_TYPE.PARAGRAPH)
        new_heading_style.base_style = styles['Heading 1']
        font = new_heading_style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)
        font.color.rgb = RGBColor(0, 0, 0)

        normalStyle = document.styles['Normal']
        normalFont = normalStyle.font
        normalFont.name = 'Courier New'
        normalFont.size = Pt(10)
    def section(document):
        sections = document.sections[0]
        sections.bottom_margin  = Mm(20)
        sections.top_margin  = Mm(20)
        sections.left_margin = Mm(20)
        sections.right_margin = Mm(20)
        section.page_height  = Mm(215.90)
        section.page_width = Mm(279.40)
    document = Document()
    styles(document)
    section(document)

    for filename in Path(argv[1]).glob('**/*.'+argv[2]):
        fr = io.open(str(filename), mode="r", encoding="utf-8")
        document.add_paragraph(filename.name, style='filename')
        document.add_paragraph(fr.read())

    document.save(argv[3])



    # fr = open(argv[1], 'r')
    # str = fr.read()
    # hash = hashlib.sha1(str.encode()).hexdigest()
    # fw = open(argv[2], 'w')
    # print(hash)
    # fw.write(hash)
